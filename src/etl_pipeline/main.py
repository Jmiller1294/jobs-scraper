import random 
import asyncio
import subprocess
from extract.scrapers.linkedin_scraper import getLinkedinJobs
from extract.file_reader import read_file
from extract.resume_router import get_resume_for_job
from transform.analyzer import analyze_job
from transform.reporter import save_results, print_summary
from load.email_digest import send_digest
from datetime import datetime
from config import CONFIG
from openai import OpenAI
from pathlib import Path
from docx import Document
from db.models import JobPosting
from db.crud import save_job
from db.session import SessionLocal
from utils import AsyncAPIKeyManager

BASE_DIR = Path(__file__).resolve().parents[1]  
DEFAULT_TEXT = """{{YOUR_NAME}}
  {{DATE}}
  
  Hiring Manager
  {{COMPANY}}
  
  Dear Hiring Manager,
  
  I am excited to apply for the {{ROLE}} position at {{COMPANY}}. With my background in [YOUR BACKGROUND], I am confident I would be a strong addition to your team.
  
  [CUSTOMIZE: 1-2 sentences about why this company specifically excites you.]
  
  In my previous roles, I have [KEY ACHIEVEMENT 1]. Additionally, I [KEY ACHIEVEMENT 2], which directly aligns with your need for [JOB REQUIREMENT].
  
  [CUSTOMIZE: Address 1-2 specific requirements from the job description with concrete examples.]
  
  I am particularly drawn to {{COMPANY}}'s mission of [COMPANY MISSION] and believe my experience in [RELEVANT SKILL] would allow me to contribute meaningfully from day one.
  
  I would welcome the opportunity to discuss how my background aligns with your team's goals. Thank you for your consideration.
  
  Sincerely,
  {{YOUR_NAME}}
"""


async def fetch_all_jobs(search_terms: list) -> list:
  all_jobs = []
  for term in search_terms:
    try:
      print(f"🔍 Searching: {term}")
      jobs = await getLinkedinJobs(search_term=term)
      if isinstance(jobs, list):
        all_jobs.extend(jobs)
      else:
        print(f"⚠️ Unexpected response for: {term}")
    except Exception as e:
      print(f"❌ Error fetching '{term}': {e}")
    await asyncio.sleep(random.uniform(2, 4))
  return all_jobs

def notify_start():
  subprocess.run([
    "osascript",
    "-e",
    'display notification "Job pipeline started" with title "Job Runner"'
  ])

def load_cover_letter_template():
  path = BASE_DIR / CONFIG["cover_letter_template_path"]
  try:
    return read_file(path)
  except FileNotFoundError:
    print("⚠️ Cover letter template not found")
    return ""

def load_cover_letter_text():
  path = BASE_DIR / CONFIG["cover_letter_text_path"]
  try:
    return read_file(path)
  except FileNotFoundError:
    print("⚠️ Cover letter text not found — using default text")
    return DEFAULT_TEXT

async def fetch_and_dedupe_jobs():
  all_jobs = await fetch_all_jobs(CONFIG["search_terms"])
  return list({
    job["job_url"]: job
    for job in all_jobs
    if job.get("job_url")
  }.values())

def create_output_dir():
  path = Path("results") / datetime.now().strftime('%Y-%m-%d_%H-%M')
  path.mkdir(parents=True, exist_ok=True)
  return path

def docx_to_text(path: str) -> str:
  doc = Document(path)
  return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

async def main(): 
  gemini_manager = AsyncAPIKeyManager(CONFIG["google_api_keys"])
  template_text = CONFIG["cover_letter_text_path"]

  print("🚀 Job pipeline started", flush=True)
  print("\n🤖 LinkedIn ATS Resume Scanner")
  print("────────────────────────────────")
  print("\n")
  notify_start()

  db = SessionLocal()

  print("Loading cover letter template.")
  print("\n")
  try:
    template_text = load_cover_letter_text()
    print("template loaded")
  except Exception as e:
    print("template not loaded")

  # ── Fetch jobs and dedupicate results────────────────────────────────────────────────────────
  jobs = await fetch_and_dedupe_jobs()

  print(f"\n📋 Total unique jobs collected: {len(jobs)}")

  existing_urls = {
    row.job_url for row in db.query(JobPosting.job_url).all()
  }
  
  # ── Initialize OpenAI and Google clients ─────────────────────────────────────────────
  open_ai_client = OpenAI(api_key=CONFIG["openai_api_key"])

  #Create file output directory
  output_dir = create_output_dir()

  # ── Analyze each job ──────────────────────────────────────────────────
  print("Analyzing job results....")
  print("\n")
  
  resume_path = CONFIG["resume_path"][0]
  results = []
  for job in jobs:
    job_url = job.get("job_url")
    # if job_url in existing_urls:
    #   continue

    resume_path = get_resume_for_job(job.get("title", ""))
    if resume_path:
      print(f"  📝 Using resume: {resume_path}")
      resume_text = docx_to_text(resume_path)
    else:
      print(f"  ⏭️  using default resume.")
      resume_text = docx_to_text(resume_path)
      
    try:
      result = await analyze_job(
        open_ai_client,
        gemini_manager,
        job,
        resume_text,
        template_text
      )

      print("saving job results")
      save_results(result, output_dir)
      save_job(job, result, db)

      result["resume_path"] = resume_path  # 👈 track per job
      results.append(result)

    except Exception as e:
      print(f"❌ {job.get('title')} @ {job.get('company')}: {e}")

  if results:
    print_summary(results)
    send_digest(results, resume_path)  # 👈 no single resume path
    print("pipeline complete!")

  db.close()







# #main function to run the pipeline
# def main():
#   db = SessionLocal()
#   print("🚀 Job started", flush=True)

#   subprocess.run([
#     "osascript",
#     "-e",
#     'display notification "Job pipeline started" with title "Job Runner"'
#   ])
  
#   print("\n🤖 LinkedIn ATS Resume Scanner")
#   print("────────────────────────────────")
#   print("\n")
#   # ── Load cover letter template ────────────────────────────────────────────────────────
#   print(f"Loading cover letter template: {BASE_DIR/CONFIG['cover_letter_template_path']}")
#   print("\n")
#   try:
#     template_text = read_file(BASE_DIR/CONFIG['cover_letter_template_path'])
#     print(f"  ✅ Template loaded ({len(template_text)} chars)")
#   except FileNotFoundError:
#     print(f"  ⚠️  Template file not found. Using built-in default template.")
#     template_text = """{{YOUR_NAME}}
#     {{DATE}}
    
#     Hiring Manager
#     {{COMPANY}}
    
#     Dear Hiring Manager,
    
#     I am excited to apply for the {{ROLE}} position at {{COMPANY}}. With my background in [YOUR BACKGROUND], I am confident I would be a strong addition to your team.
    
#     [CUSTOMIZE: 1-2 sentences about why this company specifically excites you.]
    
#     In my previous roles, I have [KEY ACHIEVEMENT 1]. Additionally, I [KEY ACHIEVEMENT 2], which directly aligns with your need for [JOB REQUIREMENT].
    
#     [CUSTOMIZE: Address 1-2 specific requirements from the job description with concrete examples.]
    
#     I am particularly drawn to {{COMPANY}}'s mission of [COMPANY MISSION] and believe my experience in [RELEVANT SKILL] would allow me to contribute meaningfully from day one.
    
#     I would welcome the opportunity to discuss how my background aligns with your team's goals. Thank you for your consideration.
    
#     Sincerely,
#     {{YOUR_NAME}}
#     """
 
#   # ── Fetch jobs ────────────────────────────────────────────────────────
#   # Run the async function and wait for results
#   all_jobs = asyncio.run(fetch_all_jobs(CONFIG["search_terms"]))

#   # Deduplicate jobs using job_url as a unique key
#   unique_jobs = {
#     job.get("job_url"): job
#     for job in all_jobs
#     if job.get("job_url")  # Only include jobs that have a URL
#   }

#   # Convert dictionary back into a list
#   final_jobs = list(unique_jobs.values())

#   # Print total unique jobs
#   print(f"\n📋 Total unique jobs collected: {len(final_jobs)}")

#   # ── Init OpenAI client ─────────────────────────────────────────────
#   open_ai_client = OpenAI(api_key=CONFIG["openai_api_key"])
#   google_gemini_client = genai.Client(api_key=CONFIG["google_api_key"])

#   # Instead of wiping the folder:
#   CONFIG["output_dir"] = f"ats_results/{datetime.now().strftime('%Y-%m-%d_%H-%M')}"
#   os.makedirs(CONFIG["output_dir"], exist_ok=True)

#   # ── Analyze each job ──────────────────────────────────────────────────
#   results = []
#   for job in final_jobs:
#     existing = db.query(JobPosting).filter_by(job_url=job.get("job_url", "<Unknown URL>")).first()
#     if existing:
#       print(f"  ⏭️  Skipping already analyzed job: {job.get('title', '<Unknown Title>')} @ {job.get('company', '<Unknown Company>')}")
#       continue

#     resume_text, resume_docx_path = get_resume_for_job(job.get("title", "<Unknown Title>"))
#     if resume_docx_path:
#       print(f"  📝 Using resume: {resume_docx_path}")
   
#     else:
#       print(f"  ⏭️  using default resume.")
#       resume_docx_path = BASE_DIR / CONFIG["resume_path"][0]  # fallback to first resume if no match
#       doc = Document(resume_docx_path)

#       text = [para.text for para in doc.paragraphs]
#       resume_text = "\n".join(text)

#       with open("output.txt", "w", encoding="utf-8") as f:
#         f.write(resume_text)

#     try:
#       result = analyze_job(open_ai_client, google_gemini_client, job, resume_text, template_text)
#       save_results(result, CONFIG["output_dir"])
#       save_job(job, result, db)
#       results.append(result)
#     except Exception as e:
#       print(f"  ❌ {job.get('title', '<Unknown Title>')} @ {job.get('company', '<Unknown Company>')}: {e}")  

#   db.close()
#   if results:
#     print_summary(results)
#     send_digest(results, resume_docx_path, updated_cover_letter_path)

if __name__ == "__main__":
  asyncio.run(main())
