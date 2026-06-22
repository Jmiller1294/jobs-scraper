import os
import smtplib
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime
from config import CONFIG
from docx.shared import Pt
from docx import Document, document
import shutil
import re
import uuid

BASE_DIR = Path(__file__).resolve().parents[3]  

def clean(s: str) -> str:
  s = s.title() or "Unknown Title"
  return re.sub(r"[^a-zA-Z0-9_ ]", "", s).replace(" ", "_")

def attach_file(msg, file_path: str):
  path = Path(file_path)
  if not path.exists():
    return

  with path.open("rb") as f:
    msg.attach(
      MIMEApplication(
        f.read(),
        Name=path.name,
        _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document"
      )
    )


def add_bullets_to_document(document: str, results: dict) -> str:
  print("adding bullets")
  bullets = results.get("improvement_bullets", [])

  # Fallback values include a short unique ID so missing data doesn't cause
  # multiple resumes to silently overwrite each other
  job_id = results.get("job_id", str(uuid.uuid4())[:8])
  company = clean(results.get("company", f"UnknownCompany_{job_id}"))
  job_title = clean(results.get("job_title", f"UnknownRole_{job_id}"))

  output_dir = f"{BASE_DIR}/{CONFIG['output_resumes_dir']}"
  os.makedirs(output_dir, exist_ok=True)  # creates it if missing, no error if it exists

  resume_filename = f"{company}_{job_title}_Resume.docx"
  resume_path = f"{output_dir}/{resume_filename}"

  try:
    shutil.copy(document, resume_path)
  except (FileNotFoundError, PermissionError) as e:
    print(f"Failed to copy resume for {company} - {job_title}: {e}") # adjust to `continue` if this is inside a loop over multiple jobs

  doc = Document(resume_path)
  """Add bullet points to a Word document."""
  
  for b in bullets:
    p = doc.add_paragraph()
    run = p.add_run(f"\n• {b}")
    run.add_break() 
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
  doc.save(f"{BASE_DIR}/{CONFIG["output_resumes_dir"]}/{company}_{job_title}_Resume.docx")
  return f"{BASE_DIR}/{CONFIG["output_resumes_dir"]}/{company}_{job_title}_Resume.docx"

def add_text_to_cover_letter_template(document: str, results: dict) -> str:
  text = results.get("cover_letter", "")
  company = clean(results.get("company", "Unknown Company"))
  job_title = clean(results.get("job_title", "Unknown Role"))

  output_dir = f"{BASE_DIR}/{CONFIG['output_cover_letters_dir']}"
  os.makedirs(output_dir, exist_ok=True)  # creates it if missing, no error if it exists

  cover_letter_filename = f"{company}_{job_title}_Cover_Letter.docx"
  cover_letter_path = f"{output_dir}/{cover_letter_filename}"

  try:
    shutil.copy(document, cover_letter_path)
  except (FileNotFoundError, PermissionError) as e:
    print(f"Failed to copy cover letter for {company} - {job_title}: {e}")
    return  "" # adjust to `continue` if this is inside a loop over multiple jobs

  doc = Document(f"{BASE_DIR}/{CONFIG["output_cover_letters_dir"]}/{company}_{job_title}_Cover_Letter.docx")

  """Add analysis rationale to the cover letter."""
  p = doc.add_paragraph()
  run = p.add_run(text)
  run.font.name = 'Garamond'
  run.font.size = Pt(10)
  doc.save(f"{BASE_DIR}/{CONFIG["output_cover_letters_dir"]}/{company}_{job_title}_Cover_Letter.docx")
  return f"{BASE_DIR}/{CONFIG["output_cover_letters_dir"]}/{company}_{job_title}_Cover_Letter.docx"

def build_html_digest(results: list, run_date: str) -> str:
  """Build the HTML email body from job results."""

  email_cfg = CONFIG["email"]
  min_score  = email_cfg["min_score"]

  # Filter and sort by score
  filtered = sorted(
    [r for r in results if r["score"] >= min_score],
    key=lambda x: x["score"],
    reverse=True
  )

  if not filtered:
    return f"""
      <html><body>
      <h2>ATS Scanner — {run_date}</h2>
      <p>No jobs scored above {min_score} today. Check back tomorrow.</p>
      </body></html>
      """

  # ── Summary bar ───────────────────────────────────────────────────
  total       = len(results)
  qualified   = len(filtered)
  avg_score   = round(sum(r["score"] for r in results) / total) if total else 0
  top_score   = filtered[0]["score"] if filtered else 0

  summary_html = f"""
  <table style="width:100%;border-collapse:collapse;margin-bottom:24px;
                background:#f8f9fa;border-radius:8px;padding:16px;">
    <tr>
      <td style="text-align:center;padding:12px;">
        <div style="font-size:28px;font-weight:700;">{total}</div>
        <div style="font-size:12px;color:#666;">Jobs Analyzed</div>
      </td>
      <td style="text-align:center;padding:12px;">
        <div style="font-size:28px;font-weight:700;color:#22c55e;">{qualified}</div>
        <div style="font-size:12px;color:#666;">Above {min_score} Score</div>
      </td>
      <td style="text-align:center;padding:12px;">
        <div style="font-size:28px;font-weight:700;">{avg_score}</div>
        <div style="font-size:12px;color:#666;">Avg Score</div>
      </td>
      <td style="text-align:center;padding:12px;">
        <div style="font-size:28px;font-weight:700;color:#3b82f6;">{top_score}</div>
        <div style="font-size:12px;color:#666;">Top Score</div>
      </td>
    </tr>
  </table>
  """

  # ── Job cards ─────────────────────────────────────────────────────
  cards_html = ""
  for r in filtered:
    score      = r["score"]
    color      = "#22c55e" if score >= 70 else "#f59e0b" if score >= 50 else "#ef4444"
    score_bar  = "█" * (score // 10) + "░" * (10 - score // 10)

    matched  = ", ".join(r.get("matched_keywords", [])[:6]) or "None"
    missing  = ", ".join(r.get("missing_keywords", [])[:6]) or "None"
    bullets  = r.get("improvement_bullets", [])
    job_url  = r.get("job_url", "#")
    strengths = r.get("strengths", [])
    weaknesses = r.get("weaknesses", [])

    bullets_html = "".join(
        f"<li style='margin-bottom:6px;'>{b}</li>"
        for b in bullets[:3]  # show top 3 bullets in email
    )

    strengths_bullets_html = "".join(
      f"<li style='margin-bottom:6px;'>{b}</li>"
      for b in strengths[:3]  # show top 3 bullets in email
    )

    weaknesses_bullets_html = "".join(
      f"<li style='margin-bottom:6px;'>{b}</li>"
      for b in weaknesses[:3]  # show top 3 bullets in email
    )

    cards_html += f"""
      <div style="border:1px solid #e5e7eb;border-radius:8px;
                  padding:20px;margin-bottom:16px;">

        <!-- Header -->
        <div style="display:flex;justify-content:space-between;
                    align-items:flex-start;margin-bottom:12px;">
          <div>
            <h3 style="margin:0;font-size:16px;">{r['job_title']}</h3>
            <div style="color:#666;font-size:14px;margin-top:2px;">
              {r['company']} · {r.get('location', '')}
            </div>
          </div>
          <div style="text-align:center;min-width:60px;">
            <div style="font-size:24px;font-weight:700;color:{color};">{score}</div>
            <div style="font-size:10px;color:#999;">ATS Score</div>
          </div>
        </div>

        <!-- Score bar -->
        <div style="font-family:monospace;font-size:12px;
                    color:{color};margin-bottom:12px;">
          [{score_bar}] {score}%
        </div>

        <!-- Keywords -->
        <table style="width:100%;margin-bottom:12px;">
          <tr>
            <td style="width:50%;vertical-align:top;padding-right:8px;">
              <div style="font-size:11px;font-weight:600;
                          color:#22c55e;margin-bottom:4px;">✅ MATCHED</div>
              <div style="font-size:12px;color:#444;">{matched}</div>
            </td>
            <td style="width:50%;vertical-align:top;padding-left:8px;">
              <div style="font-size:11px;font-weight:600;
                          color:#ef4444;margin-bottom:4px;">❌ MISSING</div>
              <div style="font-size:12px;color:#444;">{missing}</div>
            </td>
          </tr>
        </table>

        <!-- Improvement bullets -->
        <div style="background:#f8f9fa;border-radius:6px;
                    padding:12px;margin-bottom:12px;">
          <div style="font-size:11px;font-weight:600;
                      margin-bottom:6px;">🚀 TOP RESUME IMPROVEMENTS</div>
          <ul style="margin:0;padding-left:18px;font-size:12px;
                      color:#444;line-height:1.6;">
            {bullets_html}
          </ul>
        </div>

        <!-- Strengths & Weaknesses -->
        <div style="background:#f8f9fa;border-radius:6px;
                    padding:12px;margin-bottom:12px;">
          <div style="font-size:11px;font-weight:600;
                      margin-bottom:6px;">STRENGTHS</div>
          <ul style="margin:0;padding-left:18px;font-size:12px;
                      color:#444;line-height:1.6;">
            {strengths_bullets_html}
          </ul>
        </div>

        <div style="background:#f8f9fa;border-radius:6px;
                    padding:12px;margin-bottom:12px;">
          <div style="font-size:11px;font-weight:600;
                      margin-bottom:6px;">WEAKNESSES</div>
          <ul style="margin:0;padding-left:18px;font-size:12px;
                      color:#444;line-height:1.6;">
            {weaknesses_bullets_html}
          </ul>
        </div>

        <a href="{job_url}"
            style="display:inline-block;background:#1d4ed8;color:#fff;
                  padding:8px 16px;border-radius:6px;font-size:13px;
                  font-weight:500;text-decoration:none;">
          View Job →
        </a>
      </div>
    """

  # ── Full email ────────────────────────────────────────────────────
  return f"""
    <html>
    <body style="font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',
                  sans-serif;max-width:680px;margin:0 auto;padding:24px;
                  color:#111;">

      <h2 style="margin-bottom:4px;">📊 ATS Scanner Digest</h2>
      <p style="color:#666;font-size:14px;margin-bottom:24px;">{run_date}</p>

      {summary_html}

      <h3 style="margin-bottom:12px;">
        Top Matches ({qualified} jobs above {min_score})
      </h3>

      {cards_html}

      <hr style="border:none;border-top:1px solid #e5e7eb;margin:24px 0;">
      <p style="font-size:11px;color:#999;text-align:center;">
        Sent by your ATS Scanner · Full reports saved to ats_results/
      </p>

    </body>
    </html>
  """

def build_email(results: list, resume_path: str):
  print("building email...")
  email_cfg = CONFIG["email"]
  cover_letter_template_path = BASE_DIR / 'src' / CONFIG["cover_letter_template_path"]

  if not email_cfg.get("sender") or not email_cfg.get("password") or not email_cfg.get("recipient"):
    print("⚠️ Email not configured — skipping digest.")
    return None

  min_score = email_cfg["min_score"]

  # Precompute once
  qualified_jobs = [r for r in results if r["score"] >= min_score]
  qualified = len(qualified_jobs)
  top_score = max((r["score"] for r in results), default=0)

  run_date = datetime.now().strftime("%A, %B %d %Y · %I:%M %p")

  # Build message
  msg = MIMEMultipart("alternative")
  msg["Subject"] = f"ATS Digest — {qualified} matches · Top score: {top_score}/100"
  msg["From"] = email_cfg["sender"]
  msg["To"] = email_cfg["recipient"]

  html_body = build_html_digest(results, run_date)
  msg.attach(MIMEText(html_body, "html"))

  print(f" Generating and attaching files for {len(qualified_jobs)} jobs ")
  # Generate and attach files
  for job in qualified_jobs:
    new_resume = add_bullets_to_document(resume_path, job)
    new_cover_letter = add_text_to_cover_letter_template(cover_letter_template_path, job)

    attach_file(msg, new_resume)
    attach_file(msg, new_cover_letter)

  return msg
  

def send_email(msg):
  """Send the email using SMTP."""
   # Send Email
  email_cfg = CONFIG["email"]
  try:
    print("\n📧 Sending email digest...")
    with smtplib.SMTP(email_cfg["smtp_server"], email_cfg["smtp_port"]) as server:
      server.ehlo()
      server.starttls()
      server.login(email_cfg["sender"], email_cfg["password"])
      server.sendmail(
        email_cfg["sender"],
        email_cfg["recipient"],
        msg.as_string()
      )
    print(f"  ✅ Digest sent to {email_cfg['recipient']}")
  except smtplib.SMTPAuthenticationError:
    print("  ❌ Auth failed — check EMAIL_SENDER and EMAIL_PASSWORD in .env")
  except smtplib.SMTPException as e:
    print(f"  ❌ SMTP error: {e}")
  except Exception as e:
    print(f"  ❌ Unexpected error sending digest: {e}")
 

def send_digest(results: list, resume_path: str):
  """Build and send the daily digest email."""
  msg = build_email(results, resume_path)
  send_email(msg)
